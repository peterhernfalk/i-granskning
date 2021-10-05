from docx import Document
import glob

import Document_mangagement
import IS_inspection
import globals
import os
from utilities import *
#from utilities import write_output, write_output_without_newline, __extract_urls_from_table, verify_url_exists

from docx.document import Document as _Document
from docx.table import _Cell, Table, _Row
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph


NOT_FOUND = "Not found"
STYLE_FAMILY_HEADING = "Heading"
STYLE_FAMILY_RUBRIK = "Rubrik"
STYLE_FAMILY_SUBTLE_EMPHASIS = "Subtle Emphasis"
local_test = False


def __style_family(document_search_phrase):
    if "IS_*" in document_search_phrase:
        return STYLE_FAMILY_RUBRIK
    else:
        return STYLE_FAMILY_HEADING


def DOCX_prepare_inspection(document_search_phrase):
    """
    Anropar metoder för att förbereda för granskning av ett Worddokument
    """
    __set_document_name(document_search_phrase)

    if globals.IS_document_exists == True:
        __set_document_name(document_search_phrase)
        __document_structure_2_dict(__style_family(document_search_phrase))
    if globals.TKB_document_exists == True:
        __set_document_name(document_search_phrase)
        __document_structure_2_dict(__style_family(document_search_phrase))

def DOCX_inspect_revision_history(docx_document):
    """
    Kollar att dokumentets tabell med revisionshistorik har en rad för aktuell tag.
    """
    table = document.tables[1]
    antal_brister_revisionshistorik = 0
    # 2do: Display version number from page 1 in the document

    if table.cell(0,0).text == "Revisionshistorik mall":
        table = document.tables[2]

    for i, row in enumerate(table.rows):
        text = tuple(cell.text for cell in row.cells)

    if str(table.cell(i, 0).text) != globals.tag:
        write_output("OBS! Revisionshistoriken behöver uppdateras. (hittade: "+str(table.cell(i, 0).text)+" men förväntade: "+globals.tag+")")
        write_detail_box_content("<b>Resultat:</b> Revisionshistoriken behöver uppdateras. (hittade: "+str(table.cell(i, 0).text)+" men förväntade: "+globals.tag+")")
        if docx_document == globals.IS:
            #globals.IS_antal_brister_revisionshistorik = 1
            antal_brister_revisionshistorik = 1
        elif docx_document == globals.TKB:
            #globals.TKB_antal_brister_revisionshistorik = 1
            antal_brister_revisionshistorik = 1
    else:
        write_output("Revisionshistoriken är uppdaterad för denna version av domänen")
        write_detail_box_content("<b>Resultat:</b> Revisionshistoriken är uppdaterad för denna version av domänen")
    write_output("Revisionshistorikens sista rad: " + str(text))
    write_detail_box_content("Revisionshistorikens sista rad: " + str(text))
    #print(document.core_properties.author)

    # 2do: kolla om dokumentet har eftersökta custom properties
    # Custom properties finns i custom.xml i Wordfilen
    """try:
        print(document.custom_properties)
        #print("Version",document.CustomDocumentProperties('Version').value)
    except AttributeError:
        print(globals.docx_document, "has no custom properties")    # Hittills kommer alla anrop hit
    #for paragraph in document.paragraphs:
    #    print(globals.docx_document,"paragraph",paragraph.text)"""

    return antal_brister_revisionshistorik


def DOCX_display_paragraph_text_and_tables(searched_paragraph_title, display_paragraph_title, display_initial_newline, display_keylevel_text, display_tables):
    """
    Skriver ut titlar och innehåll i paragrafer och tabeller i enlighet med angivna parametrar.

    Utskrift sker via metoden: write_output

    Förbättringsförslag: lägg till en inparameter för att kunna välja om searched_paragraph_title ska matchas exakt eller ingå i hittad titel
    """
    searched_paragraph_level = DOCX_document_structure_get_levelvalue(searched_paragraph_title)
    searched_paragraph_found = False
    paragraph_or_table_found = False

    if display_paragraph_title == True and display_tables == False:
        if display_initial_newline == True:
            write_output("<br>")
            write_detail_box_html("<br>")
        paragraph_displayed = __display_paragraph_text_by_paragraph_level(searched_paragraph_level, display_keylevel_text)
        if paragraph_displayed == True:
            paragraph_or_table_found = True
    else:
        for block in __iter_block_items(document,searched_paragraph_level):
            if isinstance(block, Paragraph):
                this_paragraph_title = block.text.strip().lower()
                if this_paragraph_title == searched_paragraph_title.strip().lower():
                    searched_paragraph_found = True
                    paragraph_or_table_found = True
                    if display_paragraph_title == True:
                        __display_paragraph_text_by_paragraph_level(searched_paragraph_level, display_keylevel_text)

            elif isinstance(block, Table):
                if searched_paragraph_found == True:
                    if display_tables == True:
                        if display_paragraph_title == False:
                            write_output("<br>")
                            write_detail_box_html("<br>")
                        #__table_print(block)
                        #__table_print_beginning_columns(block)
                        __document_table_print_html_table(block)
                        paragraph_or_table_found = True
                    searched_paragraph_found = False     # Bug: supports only one table per paragraph

    return paragraph_or_table_found

def DOCX_list_searched_paragraph_titles_wrong_case(searched_paragraph_title, delimiter, searched_case):
    searched_paragraph_level = DOCX_document_structure_get_levelvalue(searched_paragraph_title)
    paragraph_title_list = []

    for block in __iter_block_items(document, searched_paragraph_level):
        if isinstance(block, Paragraph):
            this_paragraph_title = block.text.strip()
            if this_paragraph_title.lower() == searched_paragraph_title.strip().lower():

                previous_key = ""
                count = 0
                for key, value in document_paragraph_index_dict.items():
                    if key[0:len(searched_paragraph_level)] == searched_paragraph_level:
                        key_level_length = key.find(" ")
                        if len(key.strip()) > key_level_length:
                            this_key_level = key.strip()[0:key_level_length]
                            if this_key_level != previous_key:
                                key = key.replace("\n", " ")
                                key_extract = key[key.find(delimiter)+len(delimiter):]
                                if key[key.find(" "):].lower().strip() != searched_paragraph_title.lower().strip():
                                    if key[key.find(delimiter):].lower() != searched_paragraph_title.lower():
                                        if searched_case == globals.UPPER_CASE:
                                            if key_extract[0] != key_extract[0].upper():
                                                paragraph_title_list.append(key)
                                                count += 1
                                        else:
                                            if key_extract[0] != key_extract[0].lower():
                                                paragraph_title_list.append(key)
                                                count += 1
                            previous_key = key.strip()[0:key_level_length]
    return paragraph_title_list, count

def DOCX_inspect_reference_links(table_num):
    """
    Kollar om länkarna i dokumentets referenstabell fungerar eller ej.
    """
    antal_brister_referenslänkar = 0
    links = extract_urls_from_table(document, table_num)
    if len(links) == 0:
        write_output("Det finns inga länkar i referenstabellen. Obs att det ändå kan förekomma länkar med annat format (text istället för hyperlänk).")
        write_detail_box_content("Det finns inga länkar i referenstabellen. Obs att det ändå kan förekomma länkar med annat format (text istället för hyperlänk).")
    for link in links:
        status_code = verify_url_exists(link)
        if status_code == 400:
            write_output("<b>Länken är felaktig eller kan inte tolkas!</b> (statuskod: " + str(status_code) + ") för: " + link)
            write_detail_box_content("<b>Länken är felaktig eller kan inte tolkas!</b> (statuskod: " + str(status_code) + ") för: " + link)
            if globals.docx_document == globals.IS:
                #globals.IS_antal_brister_referenslänkar += 1
                antal_brister_referenslänkar += 1
            elif globals.docx_document == globals.TKB:
                #globals.TKB_antal_brister_referenslänkar += 1
                antal_brister_referenslänkar += 1
        elif status_code < 404:
            write_output("<b>OK</b> (statuskod: " + str(status_code) + ") för: <a href='" + link + "' target = '_blank'>" + link + "</a>")
            write_detail_box_content("<b>OK</b> (statuskod: " + str(status_code) + ") för: <a href='" + link + "' target = '_blank'>" + link + "</a>")
        else:
            if globals.docx_document == globals.IS:
                #globals.IS_antal_brister_referenslänkar += 1
                antal_brister_referenslänkar += 1
            elif globals.docx_document == globals.TKB:
                #globals.TKB_antal_brister_referenslänkar += 1
                antal_brister_referenslänkar += 1
            write_output("Sidan saknas! (statuskod: " + str(status_code) + ") för: " + link)
            write_detail_box_content("<b>Sidan saknas!</b> (statuskod: " + str(status_code) + ") för: " + link)

    return antal_brister_referenslänkar

def DOCX_display_paragragh_title(searched_title_name):
    """
    Söker efter angiven paragraf i lagrad dokumentstruktur. Skriver ut paragrafens titel.

    Returnerar: True om sökt paragraf hittades och False om paragrafen inte hittades
    """
    result = True
    searched_paragraph_level = DOCX_document_structure_get_exact_levelvalue(searched_title_name)
    if searched_paragraph_level != "":
        #write_output("OK. (" + searched_title_name + ") avsnitt " + searched_paragraph_level + " i TKB")
        write_output("OK. Dokument-rubrik (" + searched_paragraph_level + "):  \t" + searched_title_name)
    else:
        write_output("FEL! " + searched_title_name + " verkar inte vara beskrivet i dokumentet!")
        result = False
    return result


def __set_document_name(search_phrase):
    """
    Sätter den globala variabeln 'document' till namnet på angivet dokument
    """
    global document
    global document_name

    """os.chdir(globals.document_path)
    for word_document in glob.glob(search_phrase):
        document_name = r""+globals.document_path+"/"+word_document
    document = Document(document_name)"""
    if "IS_*" in search_phrase:
        document = globals.docx_IS_document
    elif "TKB_*" in search_phrase:
        document = globals.docx_TKB_document


def __document_structure_2_dict(style_family):
    """
    Lagrar dokumentet struktur i ett globalt dictionary.

    Lagrar även index till dokumentets paragrafer i ett globalt dictionary.
    """
    if style_family == STYLE_FAMILY_HEADING:
        level_from_style_name = {f'Heading {i}': i for i in range(10)}
    elif style_family == STYLE_FAMILY_RUBRIK:
        level_from_style_name = {f'Rubrik {i} Nr': i for i in range(10)}
    elif style_family == STYLE_FAMILY_SUBTLE_EMPHASIS:
        #2do: try other alternatives to make this style work as title
        level_from_style_name = {i for i in range(2)}
    current_levels = [0] * 10
    global document_structure_dict  #key = kapitelnamn, value = kapitelnummer
    document_structure_dict = {}
    global document_paragraph_index_dict
    document_paragraph_index_dict = {}

    index = 1
    for paragraph in document.paragraphs:
        if paragraph.style.name in level_from_style_name:
            level = level_from_style_name[paragraph.style.name]
            current_levels[level] += 1
            for level in range(level + 1, 10):
                current_levels[level] = 0
            document_structure_dict[paragraph.text.strip().lower()] = __format_levels(current_levels)
        document_paragraph_index_dict[__format_levels(current_levels) + " " + paragraph.text] = index
        index +=1

    if local_test == True:
        for value in document_structure_dict:
            print("value in document_structure_dict:",value)
        for i in document_paragraph_index_dict:
            print("i in document_paragraph_index_dict:",i)


def DOCX_document_structure_get_levelvalue(searched_key):
    """
    Söker efter angivet rubrikvärde i dictionaryt med dokumentstruktur.

    Returnerar: Om rubrikvärdet hittades så returneras dess nyckel (rubrikens titel), annars returneras NOT_FOUND
    """
    for key, value in document_structure_dict.items():
        if searched_key.strip().lower() == key:     #in key
            return value
    return NOT_FOUND

def DOCX_document_structure_get_exact_levelvalue(searched_key):
    """
    Söker efter angivet rubriktitel (nyckel) i dictionaryt med dokumentstruktur.

    Returnerar: Om rubriktiteln hittades så returneras dess rubrikvärde, annars returneras NOT_FOUND
    """
    for key, value in document_structure_dict.items():
        if searched_key.strip().lower() == key:
            return value
    return NOT_FOUND

def __display_paragraph_text_by_paragraph_level(searched_paragraph_level,display_keylevel_text):
    """
    Hämtar paragraftext från dokumentstruktur-dictionaryt med rubriknivå som nyckel, och visar den funna texten
    """
    paragraph_displayed = False
    previous_key = ""
    for key, value in document_paragraph_index_dict.items():
        if key[0:len(searched_paragraph_level)] == searched_paragraph_level:
            key_level_length = key.find(" ")
            if len(key.strip()) > key_level_length:
                this_key_level = key.strip()[0:key_level_length]
                if this_key_level == previous_key:
                    if display_keylevel_text == True:
                        write_output(globals.HTML_3_SPACES + key.strip()[key_level_length+1:])
                        write_detail_box_content(globals.HTML_3_SPACES + key.strip()[key_level_length+1:])
                        paragraph_displayed = True
                else:
                    key = key.replace("\n"," ")
                    write_output(globals.HTML_3_SPACES + key)
                    write_detail_box_content(globals.HTML_3_SPACES + key)
                    paragraph_displayed = True
                previous_key = key.strip()[0:key_level_length]
    return paragraph_displayed

def DOCX_display_paragraph_text_by_paragraph_level(searched_paragraph_level,display_keylevel_text):
    """
    Hämtar paragraftext från dokumentstruktur-dictionaryt med rubriknivå som nyckel, och visar den funna texten.

    Sökning sker efter tjänstekontraktsversion, angiven i den funna paragraftexten.
    """
    #key_level = ""
    key_text = ""
    previous_key = ""
    tk_version = ""
    for key, value in document_paragraph_index_dict.items():
        this_key_level_length = key.find(" ")
        searched_level_sublevels = searched_paragraph_level+"."
        if searched_level_sublevels in key[0:this_key_level_length]:
            key_level_length = key.find(" ")
            if len(key.strip()) > key_level_length:
                this_key_level = key.strip()[0:key_level_length]
                if this_key_level == previous_key:
                    if key_text == "version":
                        tk_version = key.strip()[key_level_length + 1:]
                else:
                    key_level = key.strip()[0:key_level_length]
                    key_text = key.lower()[key_level_length:].strip()
                previous_key = key.strip()[0:key_level_length]

    return tk_version


def __format_levels(current_level):
    """
    Formaterar angiven rubriknivå. 2do: beskriv detta bättre...

    Returnerar: formaterad rubriknivå
    """
    levels = [str(level) for level in current_level if level != 0]
    return '.'.join(levels)

def DOCX_document_structure_get_key(searched_value):
    for key, value in document_structure_dict.items():
        if searched_value == value:
            return key
    return NOT_FOUND

##### code test #####
def remove_hyperlink_tags(xml):
    import re
    #text = xml.decode('utf-8')
    text = xml
    text = text.replace("</w:hyperlink>","")
    text = re.sub('<w:hyperlink[^>]*>', "", text)
    #return text.encode('utf-8')
    return text
#####################

def __table_print(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                output_text = paragraph.text
                write_output_without_newline(globals.HTML_3_SPACES + output_text)
                write_detail_box_html(globals.HTML_3_SPACES + output_text)
        write_output("")
        write_detail_box_html("<br>")

def __table_print_beginning_columns(table):
    for row in table.rows:
        output_text = ""
        for cell in row.cells[0:3]:
            for paragraph in cell.paragraphs:
                output_text += paragraph.text + globals.HTML_3_SPACES + globals.HTML_3_SPACES + globals.HTML_3_SPACES
        write_detail_box_html(globals.HTML_3_SPACES + output_text + "<br>")


def __document_table_print_html_table(table):
    html_table = "<style> table, th, td { border:1px solid gray; empty-cells: show; } </style>"
    html_table += "<table>"
    #html_table += "<caption>"+title+"</caption>"
    row_number = 0
    for row in table.rows:
        row_number += 1
        html_table += "<tr>"
        for cell in row.cells:
            if row_number == 1:
                html_table += "<th>" + cell.text.strip() + "</th>"
            else:
                html_table += "<td>" + cell.text.strip() + "</td>"
        html_table += "</tr>"
    html_table += "</table>"
    write_detail_box_html(html_table)

"""
  <table>
    <caption>Books I May or May Not Have Read</caption>
    <tr>
      <th>Author</th>
      <th>Title</th>
      <th>Year</th>
    </tr>
    <tr>
      <td>Miguel De Cervantes</td>
      <td>The Ingenious Gentleman Don Quixote of La Mancha</td>
      <td>1605</td>
    </tr>
  </table>"""

def __iter_block_items(parent,searched_paragraph_level):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    """elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    elif isinstance(parent, _Row):
        parent_elm = parent._tr
    else:
        raise ValueError("something's not right")"""

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def DOCX_empty_table_cells_exists(table_number, display_result, display_type):
    result = False
    if globals.docx_document == globals.IS:
        globals.IS_antal_brister_tomma_tabellceller = 0
    elif globals.docx_document == globals.TKB:
        globals.TKB_antal_brister_tomma_tabellceller = 0
    antal_brister_tomma_tabellceller = 0

    html_table = ""
    if display_type == globals.DISPLAY_TYPE_TABLE:
        #html_table += "<table><caption><b><i>Tabellrader där en eller flera celler saknar innehåll</i></b></caption>"
        html_table += "<table>"

    table = document.tables[table_number]

    row_number = 0
    for row in range(0, len(table.rows)):
        row_number += 1
        column_count = len(table.row_cells(0))
        cells_missing_content = ""
        cell_contents_html = ""
        table_title = ""
        for column in range(0, column_count):
            cell_has_contents = False
            if row_number == 1:
                html_table += "<th>" + table.cell(row, column).text.strip() + "</th>"
            if table.cell(row,column).text.strip() != "":
                cell_has_contents = True
                if cell_contents_html == "":
                    cell_contents_html += "<tr>"
                cell_contents_html += "<td>" + table.cell(row,column).text + "</td>"
            else:
                if cell_contents_html == "":
                    cell_contents_html += "<tr>"
                cell_contents_html += "<td>" + table.cell(row,column).text + "&nbsp;</td>"
                for paragraph in table.cell(row, column).paragraphs:
                    xml_str = str(paragraph.paragraph_format.element.xml)
                    if "<w:t>" in xml_str or "<w:hyperlink" in xml_str or 'w:val="Hyperlink"' in  xml_str:
                        cell_has_contents = True
            if cell_has_contents == False:
                result = True
                if globals.docx_document == globals.IS:
                    table_title = IS_inspection.IS_get_infomodel_classname_from_table_number(table_number, True)
                    #globals.IS_antal_brister_tomma_tabellceller += 1
                    antal_brister_tomma_tabellceller += 1
                elif globals.docx_document == globals.TKB:
                    table_title = "TKB-tabell nummer " + str(table_number)
                    #globals.TKB_antal_brister_tomma_tabellceller += 1
                    antal_brister_tomma_tabellceller += 1
                if cells_missing_content == "":
                    cells_missing_content += str(column+1)
                else:
                    cells_missing_content += ", " + str(column+1)
                #write_detail_box_content(globals.HTML_3_SPACES + "Tabellcell utan innehåll funnen!  Tabell: " + str(table_title) + ", Rad: " + str(row) + ", Kolumn: " + str(column+1))
        cell_contents_html += "</tr>"

        if cells_missing_content != "":
            if display_type == globals.DISPLAY_TYPE_TEXT:
                write_detail_box_content(globals.HTML_3_SPACES + "Tabellceller utan innehåll!  Tabell: " + table_title + ", Rad: " + str(row) + ", Kolumn: " + cells_missing_content)
            elif display_type == globals.DISPLAY_TYPE_TABLE and result == True:
                #html_table += "<tr><td>"+table_title + "</td><td>Rad/kolumn" + str(row) + " " + cells_missing_content + "</td>" + cell_contents_html + "</tr>"
                html_table += cell_contents_html

    if display_result == True:
        if result == True:
            if display_type == globals.DISPLAY_TYPE_TABLE:
                html_table += "</table>"
                #write_detail_box_content(globals.HTML_3_SPACES + html_table)
                #print(html_table)
                write_detail_box_html(html_table)
            write_detail_box_content("<b>Resultat:</b> det finns granskade tabell(er) med en eller flera celler utan innehåll")
        else:
            write_detail_box_content("<b>Resultat:</b> alla granskade celler har innehåll")

    return result, antal_brister_tomma_tabellceller

"""def __get_infomodel_classname_from_table_number(table_number, include_level):
    #global infomodel_classes_list
    result_classtitle = ""
    for obj in IS_inspection.infomodel_classes_list:
        if obj.classtable_number == table_number:
            if include_level == True:
                result_classtitle = obj.document_level + " " + obj.classtitle
            else:
                result_classtitle = obj.classtitle
            break
    return result_classtitle"""


if local_test == True:
    TITLE = True
    NO_TITLE = False
    INITIAL_NEWLINE = True
    NO_INITIAL_NEWLINE = False
    TEXT = True
    NO_TEXT = False
    TABLES = True
    NO_TABLES = False

    #globals.document_path = "/Users/peterhernfalk/Desktop/Aktuellt/_T-granskningar/git-Repo/riv.clinicalprocess.healthcond.certificate/docs"
    #globals.document_path = "/Users/peterhernfalk/Desktop/Aktuellt/_T-granskningar/git-Repo/riv.clinicalprocess.healthcond.actoutcome/docs"
    #globals.document_path = "/Users/peterhernfalk/Desktop/Aktuellt/_T-granskningar/git-Repo/riv.clinicalprocess.activity.actions/docs"

    #print("\n*** TKB ***")
    #DOCX_prepare_inspection("TKB_*.doc*")
    #DOCX_inspect_revision_history()
    #DOCX_display_paragraph_text_and_tables("versionsinformation",TITLE,INITIAL_NEWLINE,TEXT,NO_TABLES)
    #DOCX_display_paragraph_text_and_tables("adressering",TITLE,INITIAL_NEWLINE,TEXT,NO_TABLES)
    #DOCX_display_paragraph_text_and_tables("aggregering",TITLE,INITIAL_NEWLINE,TEXT,NO_TABLES)
    #DOCX_display_paragraph_text_and_tables("sla krav",TITLE,INITIAL_NEWLINE,TEXT,TABLES)
    #DOCX_display_paragraph_text_and_tables("felhantering",TITLE,INITIAL_NEWLINE,TEXT,NO_TABLES)


    #print("\n*** IS ***")
    #DOCX_prepare_inspection("IS_*.doc*")
    #DOCX_display_paragraph_text_and_tables("klasser och attribut",TITLE,NO_INITIAL_NEWLINE,NO_TEXT,NO_TABLES)
    #DOCX_display_paragraph_text_and_tables("klasser och attribut",TITLE,INITIAL_NEWLINE,TEXT,NO_TABLES)
    #DOCX_display_paragraph_text_and_tables("Referenser",TITLE,NO_INITIAL_NEWLINE,TEXT,TABLES)
